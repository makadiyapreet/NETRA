"""
Incident report generator — PDF and DOCX output.

Generates formatted incident reports from classified posts and alerts,
suitable for law enforcement handoff and escalation workflows.

Supports:
  - PDF output via ReportLab
  - DOCX output via python-docx
  - JSON output as fallback
  - Jinja2-based escalation notice templates

Usage:
    python -m reporting.generate_report \\
        --post-ids tw-1234 ig-5678 \\
        --format pdf \\
        --output reports/incident_2024.pdf
"""

from __future__ import annotations

import json
import logging
import os
import sys
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Optional
from uuid import uuid4

logger = logging.getLogger(__name__)

PROJECT_ROOT = Path(__file__).resolve().parent.parent


class IncidentReportGenerator:
    """
    Generates incident reports from post and alert data.

    Reads from fixture files or accepts data directly.
    """

    def __init__(
        self,
        fixtures_dir: Optional[Path] = None,
    ):
        self.fixtures_dir = fixtures_dir or PROJECT_ROOT / "fixtures"

    def load_posts(self, post_ids: list[str]) -> list[dict]:
        """Load posts by ID from fixtures."""
        posts_path = self.fixtures_dir / "sample_posts.json"
        if not posts_path.exists():
            logger.warning(f"Posts file not found: {posts_path}")
            return []

        with open(posts_path) as f:
            all_posts = json.load(f)

        if not post_ids:
            return all_posts

        return [p for p in all_posts if p.get("post_id") in post_ids]

    def load_classifications(self, post_ids: list[str]) -> dict[str, dict]:
        """Load classifications by post ID from fixtures."""
        cls_path = self.fixtures_dir / "sample_classified_output.json"
        if not cls_path.exists():
            return {}

        with open(cls_path) as f:
            all_cls = json.load(f)

        result = {}
        for c in all_cls:
            if not post_ids or c.get("post_id") in post_ids:
                result[c["post_id"]] = c
        return result

    def load_alerts(self, post_ids: list[str] | None = None) -> list[dict]:
        """Load alerts from fixtures, optionally filtered by post_ids."""
        alerts_path = self.fixtures_dir / "sample_alerts_output.json"
        if not alerts_path.exists():
            return []

        with open(alerts_path) as f:
            all_alerts = json.load(f)

        if post_ids:
            return [a for a in all_alerts if a.get("post_id") in post_ids]
        return all_alerts

    def build_report_data(
        self,
        post_ids: list[str] | None = None,
        posts: list[dict] | None = None,
        classifications: dict[str, dict] | None = None,
        alerts: list[dict] | None = None,
        analyst_name: str = "NETRA System",
    ) -> dict[str, Any]:
        """
        Build the report data structure.

        Can work from either fixture files (post_ids) or direct data.
        """
        if posts is None:
            posts = self.load_posts(post_ids or [])
        if classifications is None:
            classifications = self.load_classifications(
                [p["post_id"] for p in posts]
            )
        if alerts is None:
            alerts = self.load_alerts([p["post_id"] for p in posts])

        now = datetime.now(timezone.utc)
        report_id = f"RPT-{uuid4().hex[:8].upper()}"

        # Threat breakdown
        threat_breakdown: dict[str, int] = {}
        for cls in classifications.values():
            cat = cls.get("threat_category", "Unknown")
            threat_breakdown[cat] = threat_breakdown.get(cat, 0) + 1

        # Language breakdown
        languages = set()
        for cls in classifications.values():
            languages.add(cls.get("detected_language", "unknown"))

        # Platform breakdown
        platforms = set()
        for p in posts:
            platforms.add(p.get("platform", "unknown"))

        # Date range
        timestamps = [p.get("created_at", "") for p in posts if p.get("created_at")]
        timestamps.sort()
        date_range = {
            "from": timestamps[0] if timestamps else now.isoformat(),
            "to": timestamps[-1] if timestamps else now.isoformat(),
        }

        # Engagement totals
        total_engagement = 0
        for p in posts:
            ec = p.get("engagement_counts", {})
            total_engagement += (
                ec.get("likes", 0) + ec.get("shares", 0) + ec.get("comments", 0)
            )

        # Average confidence
        confidences = [
            c.get("threat_confidence", 0) for c in classifications.values()
        ]
        avg_confidence = (
            sum(confidences) / len(confidences) if confidences else 0.0
        )

        # Build post summaries
        post_summaries = []
        for p in posts:
            cls = classifications.get(p["post_id"], {})
            post_summaries.append(
                {
                    "post_id": p["post_id"],
                    "platform": p.get("platform", "unknown"),
                    "author": p.get("author_handle", "unknown"),
                    "text_preview": p.get("text", "")[:200],
                    "threat_category": cls.get("threat_category", "Unclassified"),
                    "threat_confidence": cls.get("threat_confidence", 0),
                    "sentiment": cls.get("sentiment", "unknown"),
                    "detected_language": cls.get("detected_language", "unknown"),
                    "geo": (
                        p.get("geo_location", {}).get("place_name", "Unknown")
                        if p.get("geo_location")
                        else "Unknown"
                    ),
                    "engagement": p.get("engagement_counts", {}),
                }
            )

        # Alert summaries
        alert_summaries = []
        for a in alerts:
            alert_summaries.append(
                {
                    "alert_id": a.get("alert_id"),
                    "severity": a.get("severity"),
                    "threat_category": a.get("threat_category"),
                    "reason": a.get("triggering_reason"),
                    "bot_cluster_id": a.get("bot_cluster_id"),
                }
            )

        return {
            "report_id": report_id,
            "generated_at": now.isoformat(),
            "generated_by": analyst_name,
            "summary": {
                "total_posts": len(posts),
                "total_alerts": len(alerts),
                "threat_breakdown": threat_breakdown,
                "languages": sorted(languages),
                "platforms": sorted(platforms),
                "date_range": date_range,
                "total_engagement": total_engagement,
                "avg_confidence": round(avg_confidence, 4),
            },
            "posts": post_summaries,
            "alerts": alert_summaries,
            "recommendations": [
                "Escalate high-confidence incitement posts to law enforcement.",
                "Issue counter-narrative advisory for fake news trends.",
                "Monitor identified bot clusters for further coordinated activity.",
                "Flag author accounts for platform-level review.",
                "Update watchlist with newly identified threat keywords.",
            ],
        }

    def generate_json(self, report_data: dict, output_path: str | Path) -> Path:
        """Write report as JSON."""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w") as f:
            json.dump(report_data, f, indent=2)
        logger.info(f"JSON report written to {output_path}")
        return output_path

    def generate_pdf(self, report_data: dict, output_path: str | Path) -> Path:
        """Generate a PDF incident report using ReportLab."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import mm
            from reportlab.platypus import (
                Paragraph,
                SimpleDocTemplate,
                Spacer,
                Table,
                TableStyle,
            )
        except ImportError:
            logger.warning(
                "ReportLab not installed — falling back to JSON. "
                "Install with: pip install reportlab"
            )
            return self.generate_json(
                report_data, str(output_path).replace(".pdf", ".json")
            )

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = SimpleDocTemplate(
            str(output_path),
            pagesize=A4,
            rightMargin=20 * mm,
            leftMargin=20 * mm,
            topMargin=25 * mm,
            bottomMargin=20 * mm,
        )

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            "ReportTitle",
            parent=styles["Title"],
            fontSize=20,
            spaceAfter=10,
        )
        heading_style = ParagraphStyle(
            "ReportHeading",
            parent=styles["Heading2"],
            fontSize=14,
            spaceBefore=15,
            spaceAfter=8,
        )
        body_style = styles["BodyText"]

        elements = []

        # Title
        elements.append(Paragraph("NETRA — Incident Report", title_style))
        elements.append(
            Paragraph(
                f"Report ID: {report_data['report_id']} | "
                f"Generated: {report_data['generated_at']} | "
                f"By: {report_data['generated_by']}",
                body_style,
            )
        )
        elements.append(Spacer(1, 12))

        # Summary
        elements.append(Paragraph("Executive Summary", heading_style))
        summary = report_data["summary"]
        summary_text = (
            f"This report covers {summary['total_posts']} posts across "
            f"{', '.join(summary['platforms'])} platforms, "
            f"with {summary['total_alerts']} alerts generated. "
            f"Languages detected: {', '.join(summary['languages'])}. "
            f"Average classification confidence: {summary['avg_confidence']:.1%}. "
            f"Total engagement: {summary['total_engagement']:,}."
        )
        elements.append(Paragraph(summary_text, body_style))
        elements.append(Spacer(1, 8))

        # Threat breakdown table
        elements.append(Paragraph("Threat Breakdown", heading_style))
        threat_data = [["Category", "Count"]]
        for cat, count in summary["threat_breakdown"].items():
            threat_data.append([cat, str(count)])

        threat_table = Table(threat_data, colWidths=[120 * mm, 40 * mm])
        threat_table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e293b")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("FONTSIZE", (0, 0), (-1, -1), 10),
                    ("ALIGN", (1, 0), (1, -1), "CENTER"),
                ]
            )
        )
        elements.append(threat_table)
        elements.append(Spacer(1, 12))

        # Post details
        elements.append(Paragraph("Flagged Posts", heading_style))
        for post in report_data["posts"][:20]:  # Limit to 20 in PDF
            elements.append(
                Paragraph(
                    f"<b>{post['post_id']}</b> [{post['platform']}] "
                    f"by {post['author']} — "
                    f"<font color='red'>{post['threat_category']}</font> "
                    f"({post['threat_confidence']:.0%})",
                    body_style,
                )
            )
            elements.append(
                Paragraph(
                    f"<i>{post['text_preview'][:150]}...</i>",
                    ParagraphStyle("Preview", parent=body_style, fontSize=9),
                )
            )
            elements.append(Spacer(1, 4))

        # Recommendations
        elements.append(Paragraph("Recommendations", heading_style))
        for rec in report_data["recommendations"]:
            elements.append(Paragraph(f"• {rec}", body_style))

        doc.build(elements)
        logger.info(f"PDF report written to {output_path}")
        return output_path

    def generate_docx(self, report_data: dict, output_path: str | Path) -> Path:
        """Generate a DOCX incident report using python-docx."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.warning(
                "python-docx not installed — falling back to JSON. "
                "Install with: pip install python-docx"
            )
            return self.generate_json(
                report_data, str(output_path).replace(".docx", ".json")
            )

        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Title
        title = doc.add_heading("NETRA — Incident Report", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph(
            f"Report ID: {report_data['report_id']}  |  "
            f"Generated: {report_data['generated_at']}  |  "
            f"By: {report_data['generated_by']}"
        )
        doc.add_paragraph("")

        # Summary
        doc.add_heading("Executive Summary", level=1)
        summary = report_data["summary"]
        doc.add_paragraph(
            f"This report covers {summary['total_posts']} posts across "
            f"{', '.join(summary['platforms'])} platforms, "
            f"with {summary['total_alerts']} alerts generated. "
            f"Languages detected: {', '.join(summary['languages'])}. "
            f"Average classification confidence: {summary['avg_confidence']:.1%}. "
            f"Total engagement: {summary['total_engagement']:,}."
        )

        # Threat breakdown table
        doc.add_heading("Threat Breakdown", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Category"
        hdr[1].text = "Count"
        for cat, count in summary["threat_breakdown"].items():
            row = table.add_row().cells
            row[0].text = cat
            row[1].text = str(count)
        doc.add_paragraph("")

        # Post details
        doc.add_heading("Flagged Posts", level=1)
        for post in report_data["posts"][:30]:
            p = doc.add_paragraph()
            run = p.add_run(f"{post['post_id']} ")
            run.bold = True
            p.add_run(
                f"[{post['platform']}] by {post['author']} — "
                f"{post['threat_category']} ({post['threat_confidence']:.0%})"
            )
            preview = doc.add_paragraph(post["text_preview"][:200])
            preview.style = "Quote"

        # Alert details
        if report_data["alerts"]:
            doc.add_heading("Alerts", level=1)
            for alert in report_data["alerts"]:
                p = doc.add_paragraph()
                run = p.add_run(f"[Severity {alert['severity']}] ")
                run.bold = True
                if alert["severity"] >= 4:
                    run.font.color.rgb = RGBColor(220, 38, 38)
                p.add_run(alert.get("reason", ""))

        # Recommendations
        doc.add_heading("Recommendations", level=1)
        for rec in report_data["recommendations"]:
            doc.add_paragraph(rec, style="List Bullet")

        doc.save(str(output_path))
        logger.info(f"DOCX report written to {output_path}")
        return output_path

    def generate(
        self,
        post_ids: list[str] | None = None,
        format: str = "json",
        output_path: str | Path | None = None,
        analyst_name: str = "NETRA System",
        posts: list[dict] | None = None,
        classifications: dict[str, dict] | None = None,
        alerts: list[dict] | None = None,
    ) -> Path:
        """
        Generate an incident report in the specified format.

        Args:
            post_ids: Post IDs to include (loads from fixtures).
            format: 'json', 'pdf', or 'docx'.
            output_path: Where to write the report file.
            analyst_name: Name to credit on the report.
            posts: Direct post data (overrides fixture loading).
            classifications: Direct classification data.
            alerts: Direct alert data.

        Returns:
            Path to the generated report file.
        """
        report_data = self.build_report_data(
            post_ids=post_ids,
            posts=posts,
            classifications=classifications,
            alerts=alerts,
            analyst_name=analyst_name,
        )

        if output_path is None:
            reports_dir = PROJECT_ROOT / "reports"
            reports_dir.mkdir(exist_ok=True)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = reports_dir / f"incident_report_{ts}.{format}"

        generators = {
            "json": self.generate_json,
            "pdf": self.generate_pdf,
            "docx": self.generate_docx,
        }

        generator = generators.get(format, self.generate_json)
        return generator(report_data, output_path)


# ── CLI Entry Point ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    import argparse

    logging.basicConfig(level=logging.INFO)

    parser = argparse.ArgumentParser(description="Generate an incident report")
    parser.add_argument(
        "--post-ids", nargs="*", default=None, help="Post IDs to include"
    )
    parser.add_argument(
        "--format",
        choices=["json", "pdf", "docx"],
        default="json",
        help="Output format",
    )
    parser.add_argument(
        "--output", default=None, help="Output file path"
    )
    parser.add_argument(
        "--analyst", default="NETRA System", help="Analyst name for the report"
    )

    args = parser.parse_args()

    generator = IncidentReportGenerator()
    path = generator.generate(
        post_ids=args.post_ids,
        format=args.format,
        output_path=args.output,
        analyst_name=args.analyst,
    )
    print(f"✅ Report generated: {path}")
